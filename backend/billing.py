"""Planos Pro, Stripe (opcional) e exportação DOCX."""
from __future__ import annotations

import io
import logging
import os
from datetime import datetime
from typing import Any

logger = logging.getLogger(__name__)


def _store():
    import admin_store

    return admin_store


def billing_enabled() -> bool:
    cfg = _store().get_config()
    if cfg.get("billing_enabled") == "1":
        return True
    return os.getenv("BILLING_ENABLED", "").strip().lower() in ("1", "true", "yes")


def billing_config() -> dict[str, Any]:
    cfg = _store().get_config()
    stripe_pk = (os.getenv("STRIPE_PUBLIC_KEY") or cfg.get("stripe_public_key") or "").strip()
    price_label = cfg.get("pro_price_label") or "9,99 €/mês"
    pro_quota = int(cfg.get("pro_quota_daily") or os.getenv("PRO_QUOTA_DAILY", "200") or "200")
    return {
        "enabled": billing_enabled(),
        "stripe_configured": bool(stripe_secret_key()),
        "stripe_public_key": stripe_pk if billing_enabled() else "",
        "price_label": price_label,
        "pro_quota_daily": pro_quota,
        "checkout_ready": billing_enabled() and bool(stripe_secret_key() and stripe_price_id()),
    }


def stripe_secret_key() -> str:
    cfg = _store().get_config()
    return (os.getenv("STRIPE_SECRET_KEY") or cfg.get("stripe_secret_key") or "").strip()


def stripe_webhook_secret() -> str:
    cfg = _store().get_config()
    return (os.getenv("STRIPE_WEBHOOK_SECRET") or cfg.get("stripe_webhook_secret") or "").strip()


def stripe_price_id() -> str:
    cfg = _store().get_config()
    return (os.getenv("STRIPE_PRICE_ID_PRO") or cfg.get("stripe_price_id_pro") or "").strip()


def pro_quota_limit() -> int:
    return int(billing_config()["pro_quota_daily"])


def get_user_plan(email: str) -> str:
    email = (email or "").strip().lower()
    if not email:
        return "free"
    if not billing_enabled():
        return "free"
    conn = _store().get_connection()
    try:
        row = conn.execute(
            "SELECT plan, status FROM user_subscriptions WHERE user_email = ?",
            (email,),
        ).fetchone()
        if not row:
            return "free"
        plan = (row["plan"] or "free").lower()
        status = (row["status"] or "").lower()
        if plan == "pro" and status in ("active", "trialing", ""):
            return "pro"
        return "free"
    finally:
        conn.close()


def is_pro_user(email: str) -> bool:
    return get_user_plan(email) == "pro"


def upsert_subscription(
    email: str,
    *,
    plan: str = "pro",
    status: str = "active",
    stripe_customer_id: str | None = None,
    stripe_subscription_id: str | None = None,
    current_period_end: str | None = None,
) -> None:
    email = email.strip().lower()
    now = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    conn = _store().get_connection()
    try:
        conn.execute(
            """
            INSERT INTO user_subscriptions (
                user_email, plan, status, stripe_customer_id,
                stripe_subscription_id, current_period_end, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(user_email) DO UPDATE SET
                plan = excluded.plan,
                status = excluded.status,
                stripe_customer_id = COALESCE(excluded.stripe_customer_id, user_subscriptions.stripe_customer_id),
                stripe_subscription_id = COALESCE(excluded.stripe_subscription_id, user_subscriptions.stripe_subscription_id),
                current_period_end = excluded.current_period_end,
                updated_at = excluded.updated_at
            """,
            (
                email,
                plan,
                status,
                stripe_customer_id,
                stripe_subscription_id,
                current_period_end,
                now,
            ),
        )
        conn.commit()
    finally:
        conn.close()


def list_subscriptions(limit: int = 50) -> list[dict]:
    conn = _store().get_connection()
    try:
        rows = conn.execute(
            """
            SELECT user_email, plan, status, stripe_customer_id,
                   stripe_subscription_id, current_period_end, updated_at
            FROM user_subscriptions
            ORDER BY updated_at DESC
            LIMIT ?
            """,
            (max(1, min(limit, 200)),),
        ).fetchall()
        return [_store().row_to_dict(r) for r in rows]
    finally:
        conn.close()


def create_checkout_session(email: str, *, success_url: str, cancel_url: str) -> dict:
    if not billing_enabled():
        raise ValueError("Pagamentos desativados.")
    secret = stripe_secret_key()
    price_id = stripe_price_id()
    if not secret or not price_id:
        raise ValueError("Stripe não configurado (chave ou price ID em falta).")
    try:
        import stripe
    except ImportError as exc:
        raise ValueError("Pacote stripe não instalado no servidor.") from exc

    stripe.api_key = secret
    session = stripe.checkout.Session.create(
        mode="subscription",
        customer_email=email,
        line_items=[{"price": price_id, "quantity": 1}],
        success_url=success_url,
        cancel_url=cancel_url,
        metadata={"user_email": email},
        subscription_data={"metadata": {"user_email": email}},
    )
    return {"url": session.url, "session_id": session.id}


def handle_stripe_webhook(payload: bytes, sig_header: str) -> dict:
    secret = stripe_webhook_secret()
    if not secret:
        raise ValueError("Webhook secret em falta.")
    try:
        import stripe
    except ImportError as exc:
        raise ValueError("Pacote stripe não instalado.") from exc

    stripe.api_key = stripe_secret_key()
    event = stripe.Webhook.construct_event(payload, sig_header, secret)
    etype = event["type"]
    data = event["data"]["object"]

    if etype == "checkout.session.completed":
        email = (data.get("metadata") or {}).get("user_email") or data.get("customer_email") or ""
        if email:
            upsert_subscription(
                email,
                plan="pro",
                status="active",
                stripe_customer_id=data.get("customer"),
                stripe_subscription_id=data.get("subscription"),
            )
    elif etype in ("customer.subscription.updated", "customer.subscription.deleted"):
        meta = data.get("metadata") or {}
        email = meta.get("user_email", "")
        status = (data.get("status") or "canceled").lower()
        plan = "pro" if status in ("active", "trialing") else "free"
        if email:
            period_end = None
            if data.get("current_period_end"):
                period_end = datetime.utcfromtimestamp(int(data["current_period_end"])).isoformat() + "Z"
            upsert_subscription(
                email,
                plan=plan,
                status=status,
                stripe_customer_id=data.get("customer"),
                stripe_subscription_id=data.get("id"),
                current_period_end=period_end,
            )
    return {"ok": True, "type": etype}


def build_docx_bytes(text: str, title: str = "Transcrição Ouviescrevi") -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Exportação DOCX indisponível no servidor.") from exc

    doc = Document()
    doc.add_heading(title, level=1)
    for para in (text or "").split("\n"):
        p = para.strip()
        if p:
            doc.add_paragraph(p)
        else:
            doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
